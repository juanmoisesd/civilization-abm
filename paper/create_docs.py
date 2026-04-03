#!/usr/bin/env python3
"""
Script to create manuscript_anon.docx and title_page.docx
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    print("python-docx is available")
except ImportError:
    print("python-docx not found, installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    print("python-docx installed and imported")

OUTPUT_DIR = r"C:\Users\DELL\Pictures\Claude Industrial\civilization-abm\paper"


def add_bold_italic_run(para, text):
    """Add a run that handles **bold** and *italic* markdown inline."""
    import re
    # We'll parse simple bold/italic patterns
    # Pattern: **text** = bold, *text* = italic (non-overlapping)
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                para.add_run(part)


def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def make_bold_run(para, label, value):
    """Add a bold label followed by normal value text."""
    run = para.add_run(label)
    run.bold = True
    if value:
        para.add_run(value)


# ============================================================
# FILE 1: manuscript_anon.docx
# ============================================================

def create_manuscript_anon():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Configure heading styles
    for i, (heading_name, size) in enumerate([('Heading 1', 14), ('Heading 2', 13), ('Heading 3', 12)], 1):
        h_style = doc.styles[heading_name]
        h_style.font.name = 'Times New Roman'
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0, 0, 0)

    # Set page margins (1 inch)
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- TITLE ---
    title_para = doc.add_heading(
        'Emergent Wealth Inequality in Agent-Based Civilizations: Effects of Fiscal Policy, Network Topology, and Initial Conditions',
        level=1
    )
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- ABSTRACT heading ---
    doc.add_heading('Abstract', level=2)

    # Abstract body
    abstract_text = (
        "Agent-based modeling offers a powerful framework for studying emergent social phenomena from individual-level interactions. "
        "This paper presents Civilization-ABM, an open-source computational model of an artificial society in which heterogeneous agents "
        "interact through economic exchange, social network ties, and institutional rules. Three factors are systematically varied across "
        "11 experimental conditions: initial wealth dispersion, fiscal redistribution policy, and social network topology. Experiments use "
        "30 replications per condition over 200 simulation steps (N\u00a0=\u00a0100 agents; 66,000 total simulation runs)."
    )
    doc.add_paragraph(abstract_text)

    abstract_text2 = (
        "Four principal findings emerge. Progressive taxation reduces the Gini coefficient by 13.3% relative to the no-taxation baseline, "
        "but at a cost of 62.5% of aggregate wealth \u2014 a computational confirmation of the equity\u2013efficiency trade-off. "
        "High initial wealth dispersion (\u03c3\u00a0=\u00a01.5) paradoxically produces the lowest final inequality of all conditions tested "
        "(Gini\u00a0=\u00a00.185), 51% below the medium-dispersion baseline, through a redistribution amplification mechanism in which greater "
        "absolute wealth activates progressive tax brackets more intensively. Social network topology exerts only modest inequality effects "
        "(\u0394Gini\u00a0<\u00a00.006). A minimum wealth floor policy counterproductively increases inequality (+3.5% Gini) while reducing "
        "total wealth by 37.8%, revealing an adverse interaction with progressive redistribution."
    )
    doc.add_paragraph(abstract_text2)

    doc.add_paragraph(
        "Model code and data are available at https://github.com/jmdelaserna/civilization-abm."
    )

    # Keywords
    kw_para = doc.add_paragraph()
    run_kw = kw_para.add_run("Keywords: ")
    run_kw.bold = True
    kw_para.add_run("agent-based modeling; wealth inequality; fiscal policy; social networks; Gini coefficient; emergence")

    # ============================================================
    # INTRODUCTION
    # ============================================================
    doc.add_heading('Introduction', level=1)

    intro_paras = [
        ("The study of social inequality has occupied sociologists, economists, and political scientists for centuries, yet a fundamental "
         "challenge persists: how do macro-level patterns of wealth concentration emerge from micro-level individual decisions and interactions? "
         "Traditional analytical approaches \u2014 equilibrium models, regression analyses, survey-based methods \u2014 are powerful but structurally "
         "limited in their capacity to capture the dynamic, nonlinear, and emergent nature of social stratification (Piketty & Saez, 2014; Atkinson, 2015). "
         "Agent-based modeling (ABM) offers a complementary paradigm: rather than deriving macro-patterns from assumed equilibria, it "),
        None,  # italic segment marker
        ("grows"),
        None,
        (" them from the bottom up, allowing researchers to observe how systemic inequality emerges from the repeated interaction of "
         "heterogeneous autonomous agents (Epstein & Axtell, 1996)."),
    ]

    # Intro paragraph 1 (with italic "grows")
    p = doc.add_paragraph()
    p.add_run(
        "The study of social inequality has occupied sociologists, economists, and political scientists for centuries, yet a fundamental "
        "challenge persists: how do macro-level patterns of wealth concentration emerge from micro-level individual decisions and interactions? "
        "Traditional analytical approaches \u2014 equilibrium models, regression analyses, survey-based methods \u2014 are powerful but structurally "
        "limited in their capacity to capture the dynamic, nonlinear, and emergent nature of social stratification (Piketty & Saez, 2014; Atkinson, 2015). "
        "Agent-based modeling (ABM) offers a complementary paradigm: rather than deriving macro-patterns from assumed equilibria, it "
    )
    r_italic = p.add_run("grows")
    r_italic.italic = True
    p.add_run(
        " them from the bottom up, allowing researchers to observe how systemic inequality emerges from the repeated interaction of "
        "heterogeneous autonomous agents (Epstein & Axtell, 1996)."
    )

    doc.add_paragraph(
        "The seminal contribution of Epstein and Axtell (1996), who introduced the Sugarscape model, demonstrated that persistent and "
        "highly skewed wealth distributions could emerge spontaneously from a population of agents following elementary rules in a "
        "resource-heterogeneous environment, without any design intent or central authority. This generative approach to social science "
        "(Epstein, 2006) has since been extended to study segregation (Schelling, 1971), cooperation (Axelrod, 1984), market dynamics "
        "(Tesfatsion & Judd, 2006), epidemiological spread (Epstein, 2009), and political conflict (Cederman, 1997). However, a systematic "
        "computational study that integrates fiscal policy mechanisms, dynamic social class formation, and network-mediated interaction within "
        "a single unified artificial civilization model remains underexplored."
    )

    doc.add_paragraph(
        "The question of what drives wealth inequality is not merely academic. Global inequality has risen substantially over the past four "
        "decades (Atkinson et al., 2011), with the top 1% of households capturing an increasingly disproportionate share of total wealth in "
        "both developed and developing economies (Piketty, 2014). Piketty\u2019s empirical analysis suggests that when the rate of return on "
        "capital exceeds economic growth, wealth concentration becomes self-reinforcing. Parallel work in econophysics has demonstrated that "
        "even in the simplest closed economic systems \u2014 where agents exchange random amounts of money \u2014 the stationary distribution "
        "converges to a Boltzmann-Gibbs exponential law, implying that inequality is not an aberration but a natural thermodynamic outcome of "
        "economic exchange (Dragulescu & Yakovenko, 2000). Whether institutional interventions such as progressive taxation can meaningfully "
        "alter this trajectory is an empirical question that ABM is uniquely suited to address."
    )

    doc.add_paragraph(
        "Social network structure adds a further dimension of complexity. Real human societies exhibit the \u201csmall-world\u201d property \u2014 "
        "high local clustering combined with short global path lengths \u2014 first formalized by Watts and Strogatz (1998). They also show "
        "scale-free degree distributions characteristic of preferential attachment (Barab\u00e1si & Albert, 1999), meaning that highly connected "
        "individuals disproportionately influence resource and status flows. The topology of these networks has been shown to modulate the spread "
        "of inequality (Brzezinski & Kania, 2025), the emergence of cooperation (Nowak & May, 1992), and the resilience of social institutions "
        "(Ostrom, 1990). Yet the joint effect of network topology and fiscal policy on wealth dynamics has not been systematically examined within "
        "an ABM framework."
    )

    doc.add_paragraph(
        "This paper makes three contributions. First, it introduces Civilization-ABM, an open-source, reproducible, Mesa-based (Kazil et al., 2020) "
        "simulation platform integrating economic exchange, dynamic social class formation, reputation mechanisms, institutional rules, and "
        "NetworkX-based social graphs in a unified modular architecture. Second, it reports a systematic factorial experiment across 11 conditions, "
        "varying initial inequality, fiscal policy, and network topology, with 30 replications per condition. Third, it reports and interprets "
        "emergent macro-level outcomes \u2014 Gini coefficient, Theil index, Palma ratio, class mobility, and reputational dynamics \u2014 against "
        "predictions from econophysics, institutional economics, and complexity science."
    )

    # ============================================================
    # THEORETICAL FRAMEWORK
    # ============================================================
    doc.add_heading('Theoretical Framework and Related Work', level=1)

    doc.add_heading('Agent-Based Modeling and the Generative Approach', level=2)
    doc.add_paragraph(
        "Agent-based modeling is a computational methodology in which a system is represented as a collection of autonomous agents \u2014 each "
        "with its own attributes, rules, and memory \u2014 that interact with each other and with a shared environment (Gilbert & Troitzsch, 2005; "
        "Tesfatsion & Judd, 2006). Unlike top-down analytical models, ABM is bottom-up: macro-level phenomena emerge from the aggregate of "
        "micro-level behaviors (Epstein & Axtell, 1996). This generative epistemology, formalized by Epstein (2008) under the motto \u201cif you "
        "didn\u2019t grow it, you didn\u2019t explain it,\u201d has become a foundational principle of computational social science."
    )
    doc.add_paragraph(
        "The canonical wealth-distribution ABM dates to Sugarscape (Epstein & Axtell, 1996), in which agents harvest and accumulate a resource "
        "distributed unevenly across a landscape. Despite each agent following only simple local rules, the emergent wealth distribution closely "
        "resembles empirically observed Pareto distributions. Subsequent extensions introduced trade, disease transmission, cultural evolution "
        "(Axelrod, 1997), and combat (Epstein, 1999), demonstrating ABM\u2019s versatility as a platform for social theorizing."
    )

    doc.add_heading('Econophysics of Wealth Distribution', level=2)
    doc.add_paragraph(
        "An important theoretical benchmark is provided by the econophysics literature. Dragulescu and Yakovenko (2000) showed that in any closed "
        "economic system where agents exchange random amounts of money in pairwise transactions, the equilibrium money distribution is a "
        "Boltzmann-Gibbs exponential, where T is the \u201ceconomic temperature\u201d (mean money per agent). This result implies that without "
        "redistribution mechanisms, inequality is a thermodynamic inevitability. Models incorporating capital returns and savings propensities "
        "generate Pareto power-law tails (Yakovenko & Rosser, 2009), consistent with empirical data on the ultra-rich. These theoretical results "
        "provide the null hypothesis for our experiments: without taxation, Civilization-ABM should converge toward a highly unequal stationary distribution."
    )

    doc.add_heading('Fiscal Policy and Redistribution', level=2)
    doc.add_paragraph(
        "The impact of redistribution on inequality has been studied extensively in both empirical economics (Piketty, 2014; Milanovic, 2016) and "
        "computational models (Fagiolo & Roventini, 2017). Progressive taxation \u2014 where marginal rates increase with wealth \u2014 is theoretically "
        "predicted to compress the wealth distribution more effectively than flat-rate taxation (Atkinson et al., 2011). However, redistribution "
        "entails trade-offs: aggressive taxation may reduce incentives for wealth accumulation, potentially shrinking the total resource pool. This "
        "paper extends this line of inquiry by comparing no taxation, flat taxation, and progressive taxation within the same computational framework."
    )

    doc.add_heading('Social Networks and Inequality', level=2)
    doc.add_paragraph(
        "Watts and Strogatz (1998) demonstrated that a small amount of random rewiring in a regular network creates small-world properties. "
        "Barab\u00e1si and Albert (1999) showed that preferential attachment generates scale-free degree distributions, in which hubs may "
        "concentrate wealth through disproportionate interaction. Recent agent-based work calibrated to Italian wealth survey data confirmed that "
        "network-mediated interaction explains a substantial portion of observed wealth persistence (Brzezinski & Kania, 2025). This paper tests "
        "whether these effects persist when strong fiscal institutions are simultaneously present."
    )

    doc.add_heading('Social Class Dynamics, Mobility, and Norms', level=2)
    doc.add_paragraph(
        "Social stratification has been modeled computationally since Schelling\u2019s (1971) landmark work on residential segregation. In this "
        "model, social classes emerge endogenously as a function of each agent\u2019s wealth relative to the population mean, aligning with the "
        "sociological tradition of relative deprivation theory (Runciman, 1966) and the economic literature on intergenerational mobility "
        "(Chetty et al., 2014). Beyond economic exchange, normative mechanisms \u2014 reputation, sanctioning, ostracism \u2014 regulate behavior "
        "and sustain cooperation (Fehr & G\u00e4chter, 2002). This model incorporates reputation as an agent attribute that decays under "
        "competitive exploitation and recovers through cooperative transfers, drawing on Axelrod\u2019s (1984) evolutionary cooperation framework."
    )

    # ============================================================
    # METHODS
    # ============================================================
    doc.add_heading('Methods', level=1)

    doc.add_heading('Model Overview', level=2)
    doc.add_paragraph(
        "Civilization-ABM is implemented in Python 3.11 using the Mesa 3.x agent-based modeling framework (Kazil et al., 2020). Full source "
        "code, experimental configurations, and analysis scripts are available at https://github.com/jmdelaserna/civilization-abm under an MIT "
        "license, and the simulation model is deposited at the CoMSES Computational Model Library. The model comprises four tightly coupled "
        "modules: (1) agents, (2) social network environment, (3) institutional rules, and (4) data collection."
    )

    doc.add_heading('Agents', level=2)
    doc.add_paragraph(
        "Each simulation contains N\u00a0=\u00a0100 agents. Upon initialization, each agent is assigned: wealth (w) drawn from a log-normal "
        "distribution LN(\u00b5\u00a0=\u00a02.3, \u03c3\u00a0=\u00a0\u03c3\u2080), where \u03c3\u2080 is the initial inequality parameter; "
        "strategy (s) drawn uniformly from {cooperative, competitive, neutral}; and reputation (r) initialized at 1.0 (range: 0.0\u20132.0). "
        "Social class (c) is dynamically assigned at each step as lower (w < 0.5\u03bc\u0304), middle (0.5\u03bc\u0304 \u2264 w < 1.5\u03bc\u0304), "
        "or upper (w \u2265 1.5\u03bc\u0304), where \u03bc\u0304 is the current population mean wealth."
    )
    doc.add_paragraph(
        "At each simulation step, agents are activated in random order. Each activated agent selects a random other agent and executes a "
        "strategy-dependent wealth transfer: cooperative agents donate 5% of own wealth to poorer agents (reputation +0.02); competitive agents "
        "extract wealth from richer agents (other\u2019s reputation \u22120.05); neutral agents transfer 2.5% with probability 0.5."
    )

    doc.add_heading('Social Network', level=2)
    doc.add_paragraph(
        "Three network conditions are implemented using NetworkX (Hagberg et al., 2008): no network (fully random interaction); small-world "
        "(Watts-Strogatz graph with k\u00a0=\u00a04 and p\u00a0=\u00a00.1; Watts & Strogatz, 1998); and scale-free (Barab\u00e1si-Albert "
        "graph with m\u00a0=\u00a02; Barab\u00e1si & Albert, 1999)."
    )

    doc.add_heading('Institutional Rules', level=2)
    doc.add_paragraph(
        "Three fiscal conditions are applied at the end of each step: no taxation; flat tax (5% collected from each agent, redistributed "
        "equally); and progressive tax (marginal rates of 5% for w \u2264 20, 10% for 20 < w \u2264 50, 20% for w > 50; proceeds redistributed "
        "equally). Agents with reputation r < 0.3 incur a wealth penalty of 0.5 units per step. An optional minimum wealth floor maintains "
        "w \u2265 1.0 for all agents."
    )

    doc.add_heading('Experimental Design', level=2)
    doc.add_paragraph(
        "A single-factor design is implemented across 11 conditions (Table\u00a01), varying one independent variable at a time against a "
        "common baseline. Each condition runs for 30 independent replications with 200 simulation steps (seeds\u00a0=\u00a042, 43, \u2026, 71)."
    )

    # Table 1 caption
    p_t1 = doc.add_paragraph()
    r_t1 = p_t1.add_run("Table 1.")
    r_t1.bold = True
    p_t1.add_run(" Experimental conditions.")

    # Table 1
    t1_headers = ['Condition', '\u03c3\u2080', 'Tax policy', 'Network', 'Floor']
    t1_data = [
        ['ineq_low', '0.3', 'progressive', 'small_world', 'No'],
        ['ineq_medium', '0.8', 'progressive', 'small_world', 'No'],
        ['ineq_high', '1.5', 'progressive', 'small_world', 'No'],
        ['tax_none', '0.8', 'none', 'small_world', 'No'],
        ['tax_flat', '0.8', 'flat', 'small_world', 'No'],
        ['tax_progressive', '0.8', 'progressive', 'small_world', 'No'],
        ['net_none', '0.8', 'progressive', 'none', 'No'],
        ['net_small_world', '0.8', 'progressive', 'small_world', 'No'],
        ['net_scale_free', '0.8', 'progressive', 'scale_free', 'No'],
        ['floor_off', '0.8', 'progressive', 'small_world', 'No'],
        ['floor_on', '0.8', 'progressive', 'small_world', 'Yes'],
    ]

    table1 = doc.add_table(rows=1 + len(t1_data), cols=5)
    table1.style = 'Table Grid'
    # Header row
    hdr_cells = table1.rows[0].cells
    for i, h in enumerate(t1_headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    # Data rows
    for row_idx, row_data in enumerate(t1_data):
        row_cells = table1.rows[row_idx + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = val

    doc.add_heading('Outcome Measures', level=2)
    doc.add_paragraph(
        "The following metrics are computed at each step and averaged across 30 replications: Gini coefficient (Gini, 1921); Theil index T "
        "(Theil, 1967); Palma ratio (Palma, 2011); mean and total wealth; upper and lower class fractions; mean reputation; and network "
        "clustering coefficient."
    )

    # ============================================================
    # RESULTS
    # ============================================================
    doc.add_heading('Results', level=1)

    doc.add_heading('Baseline Dynamics', level=2)
    doc.add_paragraph(
        "A representative single simulation under baseline conditions (N\u00a0=\u00a0100, \u03c3\u2080\u00a0=\u00a00.8, progressive tax, "
        "small-world network, seed\u00a0=\u00a042) converged after 200 steps to a Gini coefficient of 0.336 and a strategy entropy of "
        "1.580 bits \u2014 approaching the theoretical maximum for three strategies (log\u2082 3\u00a0=\u00a01.585 bits), indicating sustained "
        "behavioral diversity with no dominant strategy, consistent with evolutionary game-theoretic predictions (Axelrod, 1984). The baseline "
        "Gini falls within the empirically observed range for mixed-economy European nations, providing an initial plausibility check for the "
        "model. The full systematic results are summarized in Table\u00a02."
    )

    # Table 2 caption
    p_t2 = doc.add_paragraph()
    r_t2 = p_t2.add_run("Table 2.")
    r_t2.bold = True
    p_t2.add_run(" Final-step outcome metrics averaged across 30 replications per condition.")

    # Table 2
    t2_headers = ['Condition', 'Gini', 'Mean Wealth', 'Total Wealth', 'Upper Class', 'Lower Class', 'Mean Reputation', 'Clustering']
    t2_data = [
        ['ineq_low', '0.362', '0.436', '43.61', '0.254', '0.333', '1.001', '0.382'],
        ['ineq_medium', '0.378', '0.714', '71.40', '0.291', '0.343', '0.902', '0.382'],
        ['ineq_high', '0.185', '11.888', '1188.84', '0.133', '0.150', '0.825', '0.382'],
        ['tax_none', '0.436', '1.900', '189.95', '0.256', '0.336', '0.970', '0.382'],
        ['tax_flat', '0.383', '0.690', '69.03', '0.286', '0.350', '0.893', '0.382'],
        ['tax_progressive', '0.378', '0.714', '71.40', '0.291', '0.343', '0.902', '0.382'],
        ['net_none', '0.383', '0.693', '69.25', '0.277', '0.349', '0.899', '0.000'],
        ['net_scale_free', '0.377', '0.710', '70.96', '0.298', '0.343', '0.894', '0.123'],
        ['net_small_world', '0.378', '0.714', '71.40', '0.291', '0.343', '0.902', '0.382'],
        ['floor_off', '0.378', '0.714', '71.40', '0.291', '0.343', '0.902', '0.382'],
        ['floor_on', '0.391', '0.443', '44.34', '0.316', '0.354', '0.847', '0.382'],
    ]

    table2 = doc.add_table(rows=1 + len(t2_data), cols=8)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    for i, h in enumerate(t2_headers):
        hdr_cells2[i].text = h
        hdr_cells2[i].paragraphs[0].runs[0].bold = True
    for row_idx, row_data in enumerate(t2_data):
        row_cells2 = table2.rows[row_idx + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells2[col_idx].text = val

    doc.add_heading('Effect of Fiscal Policy', level=2)
    doc.add_paragraph(
        "The no-taxation condition produced the highest Gini coefficient (0.436), consistent with econophysics predictions that unconstrained "
        "exchange converges toward highly skewed distributions (Dragulescu & Yakovenko, 2000). Progressive taxation reduced inequality to "
        "Gini\u00a0=\u00a00.378 \u2014 a 13.3% reduction \u2014 marginally outperforming flat taxation (Gini\u00a0=\u00a00.383). Both redistribution "
        "conditions entailed substantial aggregate wealth costs: total wealth under progressive taxation (71.4 units) represented only 37.6% of "
        "the no-tax baseline (189.9 units), confirming computationally the classical equity\u2013efficiency trade-off (Atkinson et al., 2011; "
        "Piketty, 2014). Figure\u00a01 shows the temporal evolution of the Gini coefficient across the three policy conditions."
    )

    doc.add_heading('Effect of Initial Wealth Dispersion \u2014 The Redistribution Amplification Paradox', level=2)
    doc.add_paragraph(
        "The most striking result emerged from the initial inequality experiment. Contrary to the prediction that greater initial inequality "
        "produces greater final inequality, the high-dispersion condition (\u03c3\u00a0=\u00a01.5) converged to the lowest Gini of all eleven "
        "conditions: 0.185 \u2014 a value 51.0% below the medium-dispersion baseline and 57.5% below the no-tax condition. Mean wealth under "
        "ineq_high (11.888 units) was 16.6 times greater than under ineq_medium (0.714 units), and total wealth 16.6-fold higher. This wealth "
        "amplification activates the progressive tax brackets more aggressively, producing much larger absolute transfers that substantially "
        "compress the distribution. We term this phenomenon the redistribution amplification effect: in models with progressive institutional "
        "redistribution, high initial wealth dispersion can paradoxically produce lower equilibrium inequality than moderate initial dispersion, "
        "because absolute wealth levels determine the intensity of redistribution independently of relative dispersion. Figure\u00a02 illustrates "
        "this paradox."
    )

    doc.add_heading('Effect of Social Network Topology', level=2)
    doc.add_paragraph(
        "Network topology exerted the smallest effect of all three experimental factors. Gini coefficients ranged from 0.377 (scale-free) to "
        "0.383 (no network) \u2014 a spread of only 0.006, representing a 1.6% difference. The no-network condition produced the highest "
        "inequality, indicating that network-mediated interaction consistently reduces inequality regardless of topology. Scale-free networks "
        "produced marginally lower inequality than small-world networks, contrary to theoretical predictions that hub-dominated networks "
        "concentrate wealth (Barab\u00e1si & Albert, 1999; Brzezinski & Kania, 2025). This may be because hub agents participate more frequently "
        "in cooperative transfers, partially offsetting hub-mediated concentration."
    )

    doc.add_heading('Effect of Minimum Wealth Floor \u2014 The Floor Policy Backfire', level=2)
    doc.add_paragraph(
        "The minimum wealth floor policy produced results opposite to its redistributive intent. Under floor_on, the Gini coefficient increased "
        "from 0.378 to 0.391 (+3.5%), total wealth fell by 37.8%, and mean reputation declined to 0.847 \u2014 the lowest value across all "
        "conditions. The floor simultaneously increased both upper-class and lower-class fractions, shrinking the middle class. This "
        "\u201cpolarization\u201d pattern suggests that the floor mechanism disproportionately depletes middle-wealth agents, who contribute to "
        "floor maintenance but receive proportionally less redistribution benefit. Figure\u00a03 illustrates the polarization effect."
    )

    # ============================================================
    # DISCUSSION
    # ============================================================
    doc.add_heading('Discussion', level=1)

    doc.add_heading('Confirmation and Extension of Theoretical Predictions', level=2)
    doc.add_paragraph(
        "The baseline finding \u2014 unconstrained exchange produces a highly unequal distribution (Gini\u00a0=\u00a00.436) \u2014 is fully "
        "consistent with Dragulescu and Yakovenko\u2019s (2000) statistical mechanics prediction. Progressive taxation shifts the distribution "
        "toward greater equality, as predicted by redistribution theory (Piketty, 2014; Milanovic, 2016), but at a substantial aggregate cost "
        "\u2014 also as theoretically anticipated (Atkinson et al., 2011). The modest effect of network topology (\u0394Gini\u00a0=\u00a00.006) "
        "contrasts with studies reporting stronger network effects (Brzezinski & Kania, 2025; Nowak & May, 1992). This discrepancy may reflect a "
        "crowding-out effect: robust fiscal redistribution mechanisms may overshadow structural advantages conferred by network position. This "
        "suggests that network-based inequality interventions may be more effective in low-institution contexts than in environments with strong "
        "redistribution policies."
    )

    doc.add_heading('The Redistribution Amplification Paradox', level=2)
    doc.add_paragraph(
        "The redistribution amplification effect \u2014 high initial dispersion producing lowest final inequality \u2014 has no precedent in "
        "standard equilibrium models of redistribution, where initial inequality typically propagates to final inequality unless strong corrective "
        "policies are applied. The mechanism is emergent: it arises from the interaction between log-normal wealth initialization, the progressive "
        "marginal tax schedule, and equal redistribution. The same institutional rule becomes far more powerful when operating on a wealth-rich "
        "environment. This finding resonates with an apparent paradox in development economics: societies with high initial wealth inequality but "
        "strong institutions sometimes achieve better distributional outcomes than those with low inequality but weak institutions (Milanovic, 2016). "
        "Our model offers a computational mechanism for this pattern."
    )

    doc.add_heading('The Floor Policy Backfire Effect', level=2)
    doc.add_paragraph(
        "The counterproductive effect of the minimum wealth floor illustrates the general principle that institutional mechanisms can interact "
        "adversely when implemented simultaneously (Ostrom, 1990). The floor creates a \u201ctransfer trap\u201d for middle-wealth agents: they "
        "contribute to floor maintenance but do not receive disproportionate benefit. This finding is relevant to debates about universal basic "
        "income and unconditional transfer programs: wealth floor policies combined with progressive taxation may create adverse interaction effects "
        "that are not visible when either mechanism is analyzed in isolation."
    )

    doc.add_heading('Limitations and Future Directions', level=2)
    doc.add_paragraph(
        "Several limitations warrant acknowledgment. The model population is small (N\u00a0=\u00a0100) and the simulation horizon short "
        "(200 steps), limiting exploration of long-run dynamics. Agent strategies are fixed and do not evolve; a natural extension would implement "
        "evolutionary dynamics (Axelrod, 1984; Nowak & May, 1992). The network is quasi-static; real social networks are far more dynamic, with "
        "tie formation driven by wealth similarity, reputation, and proximity. The model also excludes capital returns, credit, and inheritance "
        "\u2014 mechanisms identified as primary drivers of long-run wealth concentration (Piketty, 2014). Introducing spatial heterogeneity "
        "would allow examination of regional economic divergence (Chetty et al., 2014)."
    )

    # ============================================================
    # CONCLUSION
    # ============================================================
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        "This paper presents Civilization-ABM, a systematic agent-based investigation of how initial conditions, fiscal institutions, and social "
        "network topology jointly shape the emergence of wealth inequality. Across 11 experimental conditions and 330 independent replications "
        "(66,000 total simulation runs), four principal findings emerge."
    )
    doc.add_paragraph(
        "Progressive redistribution reduces inequality at an efficiency cost. Progressive taxation achieved the lowest Gini among policy conditions "
        "(0.378), outperforming flat taxation (0.383) and no taxation (0.436), but at a 62.4% reduction in aggregate wealth, confirming the "
        "equity\u2013efficiency trade-off as a structurally robust feature of redistribution models."
    )
    doc.add_paragraph(
        "High initial wealth dispersion paradoxically minimizes final inequality. The high-dispersion condition (\u03c3\u00a0=\u00a01.5) produced "
        "a final Gini of 0.185 \u2014 51% below the medium-dispersion baseline \u2014 through a redistribution amplification mechanism in which "
        "greater absolute wealth intensifies progressive tax extraction. This emergent result challenges the assumption that initial inequality "
        "predicts final inequality."
    )
    doc.add_paragraph(
        "Network topology is a secondary determinant of inequality. Gini variation across topology conditions was less than 0.006, suggesting "
        "that robust fiscal institutions overshadow network-structural effects."
    )
    doc.add_paragraph(
        "Floor policies can backfire when combined with progressive redistribution. Minimum wealth guarantees increased the Gini coefficient and "
        "reduced total wealth by hollowing out the middle class, underscoring the importance of modeling institutional combinations rather than "
        "individual policies in isolation."
    )
    doc.add_paragraph(
        "All model code, data, and experimental configurations are openly available at https://github.com/jmdelaserna/civilization-abm and at "
        "the CoMSES Computational Model Library."
    )

    # ============================================================
    # ACKNOWLEDGEMENTS
    # ============================================================
    doc.add_heading('Acknowledgements', level=1)
    doc.add_paragraph(
        "The author used Claude (Anthropic) for assistance with manuscript drafting and editorial revision. All scientific concepts, research "
        "design, computational model implementation, experimental execution, data analysis, and intellectual conclusions are entirely the "
        "author\u2019s own."
    )

    # ============================================================
    # AUTHOR CONTRIBUTIONS
    # ============================================================
    doc.add_heading('Author Contributions', level=1)
    doc.add_paragraph(
        "The author: Conceptualization, methodology, software, formal analysis, data curation, writing \u2014 original draft, "
        "writing \u2014 review and editing, visualization."
    )

    # ============================================================
    # STATEMENTS AND DECLARATIONS
    # ============================================================
    doc.add_heading('Statements and Declarations', level=1)

    def add_bold_label_para(doc, label, text):
        p = doc.add_paragraph()
        r = p.add_run(label + " ")
        r.bold = True
        p.add_run(text)

    add_bold_label_para(doc, "Ethical considerations:", "This study uses only computational simulation data. No human participants, human data, human tissue, or personal information of any kind were involved. Ethical approval was not required.")
    add_bold_label_para(doc, "Consent to participate:", "Not applicable. This study does not involve human participants.")
    add_bold_label_para(doc, "Consent for publication:", "Not applicable. This study does not contain data from any individual person.")
    add_bold_label_para(doc, "Declaration of conflicting interest:", "The author declared no potential conflicts of interest with respect to the research, authorship, and/or publication of this article.")
    add_bold_label_para(doc, "Funding:", "This research received no external funding. It was conducted independently using open-source computational tools.")
    add_bold_label_para(doc, "Data availability:", "All simulation code, experimental configurations, raw data, and figure-generation scripts are openly available at https://github.com/jmdelaserna/civilization-abm (MIT license) and at the CoMSES Computational Model Library. The dataset supporting the findings of this study (all_conditions.csv) is included in the repository.")

    # ============================================================
    # REFERENCES
    # ============================================================
    doc.add_heading('References', level=1)

    references = [
        "Atkinson, A. B. (2015). Inequality: What can be done? Harvard University Press.",
        "Atkinson, A. B., Piketty, T., & Saez, E. (2011). Top incomes in the long run of history. Journal of Economic Literature, 49(1), 3\u201371.",
        "Axelrod, R. (1984). The evolution of cooperation. Basic Books.",
        "Axelrod, R. (1997). The dissemination of culture: A model with local convergence and global polarization. Journal of Conflict Resolution, 41(2), 203\u2013226.",
        "Barab\u00e1si, A. L., & Albert, R. (1999). Emergence of scaling in random networks. Science, 286(5439), 509\u2013512.",
        "Brzezinski, M., & Kania, A. (2025). Persistence of wealth inequality from network effects. PLOS Complex Systems. https://doi.org/10.1371/journal.pcsy.0000050",
        "Cederman, L. E. (1997). Emergent actors in world politics. Princeton University Press.",
        "Chetty, R., Hendren, N., Kline, P., & Saez, E. (2014). Where is the land of opportunity? The geography of intergenerational mobility in the United States. Quarterly Journal of Economics, 129(4), 1553\u20131623.",
        "Dragulescu, A., & Yakovenko, V. M. (2000). Statistical mechanics of money. European Physical Journal B, 17(4), 723\u2013729.",
        "Epstein, J. M. (1999). Agent-based computational models and generative social science. Complexity, 4(5), 41\u201360.",
        "Epstein, J. M. (2006). Generative social science: Studies in agent-based computational modeling. Princeton University Press.",
        "Epstein, J. M. (2008). Why model? Journal of Artificial Societies and Social Simulation, 11(4), 12. https://www.jasss.org/11/4/12.html",
        "Epstein, J. M. (2009). Modeling to contain pandemics. Nature, 460(7256), 687.",
        "Epstein, J. M., & Axtell, R. (1996). Growing artificial societies: Social science from the bottom up. MIT Press.",
        "Fagiolo, G., & Roventini, A. (2017). Macroeconomic policy in DSGE and agent-based models Redux. Journal of Artificial Societies and Social Simulation, 20(1), 1. https://www.jasss.org/20/1/1.html",
        "Fehr, E., & G\u00e4chter, S. (2002). Altruistic punishment in humans. Nature, 415(6868), 137\u2013140.",
        "Gilbert, N., & Troitzsch, K. G. (2005). Simulation for the social scientist (2nd ed.). Open University Press.",
        "Gini, C. (1921). Measurement of inequality of incomes. Economic Journal, 31(121), 124\u2013126.",
        "Grimm, V., Berger, U., Bastiansen, F., Eliassen, S., Ginot, V., Giske, J., & DeAngelis, D. L. (2006). A standard protocol for describing individual-based and agent-based models. Ecological Modelling, 198(1\u20132), 115\u2013126.",
        "Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network structure, dynamics, and function using NetworkX. In Proceedings of the 7th Python in science conference (SciPy2008) (pp. 11\u201315).",
        "Kazil, J., Masad, D., & Crooks, A. (2020). Utilizing Python for agent-based modeling: The Mesa framework. In Social, cultural, and behavioral modeling (SBP-BRiMS 2020), Lecture Notes in Computer Science vol. 12268 (pp. 308\u2013317). Springer.",
        "Milanovic, B. (2016). Global inequality: A new approach for the age of globalization. Harvard University Press.",
        "Nowak, M. A., & May, R. M. (1992). Evolutionary games and spatial chaos. Nature, 359(6398), 826\u2013829.",
        "Ostrom, E. (1990). Governing the commons: The evolution of institutions for collective action. Cambridge University Press.",
        "Palma, J. G. (2011). Homogeneous middles vs. heterogeneous tails, and the end of the \u201cInverted-U.\u201d Development and Change, 42(1), 87\u2013153.",
        "Piketty, T. (2014). Capital in the twenty-first century. Harvard University Press.",
        "Piketty, T., & Saez, E. (2014). Inequality in the long run. Science, 344(6186), 838\u2013843.",
        "Runciman, W. G. (1966). Relative deprivation and social justice. Routledge.",
        "Schelling, T. C. (1971). Dynamic models of segregation. Journal of Mathematical Sociology, 1(2), 143\u2013186.",
        "Tesfatsion, L., & Judd, K. L. (Eds.). (2006). Handbook of computational economics: Agent-based computational economics (Vol. 2). Elsevier.",
        "Theil, H. (1967). Economics and information theory. North-Holland.",
        "Watts, D. J., & Strogatz, S. H. (1998). Collective dynamics of \u201csmall-world\u201d networks. Nature, 393(6684), 440\u2013442.",
        "Yakovenko, V. M., & Rosser, J. B. (2009). Colloquium: Statistical mechanics of money, wealth, and income. Reviews of Modern Physics, 81(4), 1703\u20131725.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "manuscript_anon.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ============================================================
# FILE 2: title_page.docx
# ============================================================

def create_title_page():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Set page margins
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_field(doc, label, value):
        p = doc.add_paragraph()
        r_label = p.add_run(label + " ")
        r_label.bold = True
        p.add_run(value)
        return p

    add_field(doc, "Title:", "Emergent Wealth Inequality in Agent-Based Civilizations: Effects of Fiscal Policy, Network Topology, and Initial Conditions")
    doc.add_paragraph()  # blank line

    add_field(doc, "Author:", "Juan Mois\u00e9s de la Serna Tuya")
    add_field(doc, "Affiliation:", "Universidad Internacional de La Rioja (UNIR), Spain")
    add_field(doc, "Postal address:", "Av. de la Paz, 137, 26006 Logro\u00f1o, La Rioja, Spain")
    add_field(doc, "Phone:", "[to be filled by author]")
    add_field(doc, "Email:", "juanmoises.delaserna@unir.net")
    add_field(doc, "ORCID:", "0000-0002-8401-8018")
    doc.add_paragraph()  # blank line

    add_field(doc, "Acknowledgements:", "The author used Claude (Anthropic) for assistance with manuscript drafting and editorial revision. All scientific concepts, research design, computational model implementation, experimental execution, data analysis, and intellectual conclusions are entirely the author\u2019s own.")
    doc.add_paragraph()

    add_field(doc, "Declaration of conflicting interest:", "The author declared no potential conflicts of interest with respect to the research, authorship, and/or publication of this article.")
    doc.add_paragraph()

    add_field(doc, "Funding statement:", "This research received no external funding.")
    doc.add_paragraph()

    add_field(doc, "Ethical approval:", "Not required. This study uses only computational simulation data. No human participants, human data, or human tissue were involved.")
    doc.add_paragraph()

    add_field(doc, "Consent to participate:", "Not applicable.")
    doc.add_paragraph()

    add_field(doc, "Consent for publication:", "Not applicable.")
    doc.add_paragraph()

    add_field(doc, "Data availability:", "All simulation code, experimental configurations, raw data, and figure-generation scripts are openly available at https://github.com/jmdelaserna/civilization-abm (MIT license) and at the CoMSES Computational Model Library.")

    out_path = os.path.join(OUTPUT_DIR, "title_page.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    print("Creating manuscript_anon.docx...")
    p1 = create_manuscript_anon()
    print("Creating title_page.docx...")
    p2 = create_title_page()
    print("\nDone!")
    # Verify file sizes
    import os
    for p in [p1, p2]:
        size = os.path.getsize(p)
        print(f"  {os.path.basename(p)}: {size:,} bytes")
